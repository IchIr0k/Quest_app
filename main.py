import os
import io
import shutil
import uuid
from typing import Optional, List
from datetime import datetime

import urllib.parse

from fastapi import (
    FastAPI, Request, Form, UploadFile, File, Depends, HTTPException
)
from fastapi.responses import HTMLResponse, RedirectResponse, JSONResponse, FileResponse, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from starlette.middleware.sessions import SessionMiddleware
from sqlalchemy.orm import Session
from sqlalchemy.exc import IntegrityError

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches
from docx.oxml.shared import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill, Border, Side
from openpyxl.utils import get_column_letter
from openpyxl.drawing.image import Image as XLImage

from reportlab.lib.pagesizes import A4
from reportlab.pdfgen import canvas
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib import colors
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

import json
from datetime import datetime

from database import engine, Base, SessionLocal
import models
import crud
from auth import hash_password, verify_password, get_db, get_current_user, require_admin
from schemas import QuestCreate
import uvicorn

# --- Подготовка директорий ---
os.makedirs("static/uploads", exist_ok=True)
os.makedirs("static/images", exist_ok=True)
os.makedirs("templates", exist_ok=True)

app = FastAPI()
app.add_middleware(SessionMiddleware, secret_key="!secret_dev_change_me!")

# --- Статика и шаблоны ---
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# --- Создание таблиц ---
Base.metadata.create_all(bind=engine)


# --- Создание дефолтного админа ---
def create_default_admin():
    db = SessionLocal()
    try:
        admin = db.query(models.User).filter_by(username="admin").first()
        if not admin:
            a = models.User(
                username="admin",
                email="admin@example.com",
                hashed_password=hash_password("admin"),
                is_admin=True
            )
            db.add(a)
            db.commit()
            print("✅ Created default admin (username=admin password=admin). Change password immediately.")
    finally:
        db.close()


create_default_admin()


# --- Регистрация шрифтов для PDF ---
def register_fonts():
    try:
        # Попробуем использовать стандартные шрифты
        pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
        pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
        return True
    except:
        try:
            # Альтернативные шрифты
            pdfmetrics.registerFont(TTFont('Helvetica', 'helvetica.ttf'))
            pdfmetrics.registerFont(TTFont('Helvetica-Bold', 'helveticab.ttf'))
            return True
        except:
            return False


register_fonts()


# --- Хелперы ---
def save_upload(file: UploadFile) -> str:
    """Сохраняет файл в static/uploads и возвращает относительный путь"""
    ext = os.path.splitext(file.filename)[1]
    safe_name = f"{uuid.uuid4().hex}{ext}"
    dest = os.path.join("static", "uploads", safe_name)
    with open(dest, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)
    return f"uploads/{safe_name}"


def create_statement_template():
    """Создает шаблон заявления если его нет"""
    template_path = "templates/statement_template.docx"

    doc = Document()

    # Заголовок
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("Заявление об отказе от претензий")
    title_run.bold = True
    title_run.font.size = Pt(14)

    doc.add_paragraph()  # Пустая строка

    # Текст с метками
    content = doc.add_paragraph()
    content.add_run("Я, {full_name},")
    content.add_run("\n(Ф.И.О.)\n\n")
    content.add_run("серия {passport_series} номер {passport_number} паспорта,\n\n")

    content.add_run(
        'будучи зарегистрированным пользователем системы бронирования квестов "Алиби" и сознавая степень риска и особенности, связанные с участием в квестах с актерами, добровольно заявляю о своем решении принять участие в данном виде развлечений.\n\n')

    content.add_run('Я полностью осознаю и добровольно принимаю на себя все риски, связанные с:\n')
    content.add_run('• психологическим воздействием и элементами страха в ходе прохождения квеста;\n')
    content.add_run('• физической активностью и перемещениями в условиях ограниченного пространства;\n')
    content.add_run('• взаимодействием с актерами и импровизационными элементами программы;\n')
    content.add_run('• нахождением в помещениях со специальными эффектами (световые, звуковые, дымовые и др.).\n\n')

    content.add_run('Я подтверждаю, что:\n')
    content.add_run('• не имею медицинских противопоказаний к участию в активных играх;\n')
    content.add_run('• не страдаю сердечно-сосудистыми заболеваниями;\n')
    content.add_run('• не имею психических расстройств;\n')
    content.add_run('• не нахожусь в состоянии алкогольного или наркотического опьянения;\n')
    content.add_run('• предупрежден о возможности фото- и видеосъемки в ходе квеста.\n\n')

    content.add_run('С условиями участия ознакомлен и согласен.\n\n')

    # Подпись
    sign = doc.add_paragraph()
    sign.add_run("Ф.И.О. участника: _________________________")
    sign.add_run("\n\n(подпись)\n\n")
    sign.add_run("Дата: {current_date}.")

    doc.save(template_path)
    return template_path


# --- Маршруты ---
@app.get("/", response_class=HTMLResponse)
def index(request: Request, q: Optional[str] = None, genre: Optional[str] = None,
          difficulty: Optional[str] = None, sort: Optional[str] = None,
          skip: int = 0, db: Session = Depends(get_db)):
    filters = {}
    if q:
        filters["q"] = q
    if genre:
        filters["genre"] = genre
    if difficulty:
        filters["difficulty"] = difficulty
    if sort:
        filters["sort"] = sort

    quests = crud.get_quests(db, skip=skip, limit=15, filters=filters)
    try:
        user = get_current_user(request, db)
    except:
        user = None
    return templates.TemplateResponse("index.html", {
        "request": request,
        "quests": quests,
        "user": user,
        "skip": skip,
        "now": datetime.now
    })


@app.get("/api/quests", response_class=HTMLResponse)
def api_quests(request: Request, q: Optional[str] = None, genre: Optional[str] = None,
               difficulty: Optional[str] = None, skip: int = 0, limit: int = 6, db: Session = Depends(get_db)):
    filters = {}
    if q:
        filters["q"] = q
    if genre:
        filters["genre"] = genre
    if difficulty:
        filters["difficulty"] = difficulty
    quests = crud.get_quests(db, skip=skip, limit=15, filters=filters)
    return templates.TemplateResponse("_quest_cards.html", {"request": request, "quests": quests})


@app.get("/quest/{quest_id}", response_class=HTMLResponse)
def quest_detail(request: Request, quest_id: int, db: Session = Depends(get_db)):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(status_code=404, detail="Quest not found")

    # Получаем занятые слоты для этого квеста
    booked_slots = crud.get_booked_slots_for_date(db, quest_id, datetime.now().strftime('%Y-%m-%d'))

    try:
        user = get_current_user(request, db)
    except:
        user = None

    return templates.TemplateResponse("quest_detail.html", {
        "request": request,
        "quest": quest,
        "user": user,
        "booked_slots": booked_slots,
        "now": datetime.now
    })


@app.get("/api/available-slots")
def get_available_slots(quest_id: int, date: str, db: Session = Depends(get_db)):
    """API для получения занятых слотов"""
    booked_slots = crud.get_booked_slots_for_date(db, quest_id, date)
    return JSONResponse(booked_slots)


@app.post("/book")
def book(request: Request, quest_id: int = Form(...), date: str = Form(...), timeslot: str = Form(...),
         db: Session = Depends(get_db)):
    user = get_current_user(request, db)
    booking = crud.create_booking(db, user_id=user.id, quest_id=quest_id, date=date, timeslot=timeslot)
    if not booking:
        return JSONResponse({"success": False, "message": "Выбранный слот уже занят"}, status_code=400)
    return JSONResponse({"success": True, "message": "Бронь успешно создана"})


@app.get("/my-bookings", response_class=HTMLResponse)
def my_bookings(request: Request, db: Session = Depends(get_db)):
    """Страница с бронированиями пользователя"""
    user = get_current_user(request, db)
    bookings = crud.get_user_bookings(db, user.id)
    return templates.TemplateResponse("my_bookings.html", {
        "request": request,
        "user": user,
        "bookings": bookings,
        "now": datetime.now
    })


# --- Auth routes ---
@app.get("/login", response_class=HTMLResponse)
def login_get(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})


@app.post("/login")
def login_post(request: Request, username: str = Form(...), password: str = Form(...), db: Session = Depends(get_db)):
    user = db.query(models.User).filter_by(username=username).first()
    if not user or not verify_password(password, user.hashed_password):
        return templates.TemplateResponse("login.html", {"request": request, "error": "Неверные учётные данные"})
    request.session["user_id"] = user.id
    return RedirectResponse("/", status_code=303)


@app.get("/register", response_class=HTMLResponse)
def register_get(request: Request):
    return templates.TemplateResponse("register.html", {"request": request})


@app.post("/register")
def register_post(request: Request, username: str = Form(...), email: str = Form(None), password: str = Form(...),
                  db: Session = Depends(get_db)):
    # Проверяем, не существует ли пользователь с таким username
    existing_username = db.query(models.User).filter_by(username=username).first()
    if existing_username:
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": "Пользователь с таким именем уже существует"
        })

    # Если email указан, проверяем его уникальность
    if email:
        existing_email = db.query(models.User).filter_by(email=email).first()
        if existing_email:
            return templates.TemplateResponse("register.html", {
                "request": request,
                "error": "Пользователь с таким email уже существует"
            })

    # Проверяем длину пароля
    if len(password) < 4:
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": "Пароль должен содержать минимум 4 символа"
        })

    # Проверяем длину имени пользователя
    if len(username) < 3:
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": "Имя пользователя должно содержать минимум 3 символа"
        })

    try:
        u = models.User(
            username=username,
            email=email if email else None,
            hashed_password=hash_password(password),
            is_admin=False
        )
        db.add(u)
        db.commit()
        db.refresh(u)
        request.session["user_id"] = u.id
        return RedirectResponse("/", status_code=303)

    except IntegrityError:
        db.rollback()
        return templates.TemplateResponse("register.html", {
            "request": request,
            "error": "Произошла ошибка при создании пользователя. Попробуйте другое имя или email."
        })


@app.get("/logout")
def logout(request: Request):
    request.session.clear()
    return RedirectResponse("/", status_code=303)


# --- Admin routes ---
@app.get("/admin", response_class=HTMLResponse)
def admin_dashboard(request: Request, db: Session = Depends(get_db), user=Depends(require_admin)):
    quests = crud.get_quests(db, skip=0, limit=1000, filters={})
    return templates.TemplateResponse("admin_dashboard.html", {
        "request": request,
        "quests": quests,
        "user": user,
        "now": datetime.now
    })


@app.get("/admin/add", response_class=HTMLResponse)
def add_get(request: Request, user=Depends(require_admin)):
    return templates.TemplateResponse("add_quest.html", {"request": request, "user": user})


@app.post("/admin/add")
def add_post(request: Request,
             title: str = Form(...),
             description: str = Form(""),
             organizer_email: str = Form("alibi@mail.ru"),
             price: int = Form(2000),
             genres: List[str] = Form(...),
             difficulty: str = Form(""),
             fear_level: int = Form(0),
             players: int = Form(1),
             image: Optional[UploadFile] = File(None),
             clipboard_image: str = Form(None),
             db: Session = Depends(get_db),
             user=Depends(require_admin)):
    image_path = None

    # Обработка изображения из буфера обмена
    if clipboard_image and clipboard_image.startswith('data:image'):
        import base64
        # Извлекаем данные из Data URL
        image_data = clipboard_image.split(',')[1]
        image_bytes = base64.b64decode(image_data)

        # Сохраняем файл
        ext = '.png'  # По умолчанию PNG
        if 'image/jpeg' in clipboard_image:
            ext = '.jpg'
        elif 'image/png' in clipboard_image:
            ext = '.png'
        elif 'image/gif' in clipboard_image:
            ext = '.gif'

        safe_name = f"{uuid.uuid4().hex}{ext}"
        dest = os.path.join("static", "uploads", safe_name)

        with open(dest, "wb") as f:
            f.write(image_bytes)

        image_path = f"uploads/{safe_name}"

    # Обработка обычной загрузки файла
    elif image and image.filename:
        image_path = save_upload(image)

    genre_str = ", ".join(genres)

    new_quest = models.Quest(
        title=title,
        description=description,
        organizer_email=organizer_email,
        price=price,
        genre=genre_str,
        difficulty=difficulty,
        fear_level=fear_level,
        players=players,
        image_path=image_path
    )

    db.add(new_quest)
    db.commit()
    db.refresh(new_quest)

    return RedirectResponse("/admin", status_code=303)


@app.get("/admin/edit/{quest_id}", response_class=HTMLResponse)
def edit_get(request: Request, quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(404, "Квест не найден")
    return templates.TemplateResponse("edit_quest.html", {"request": request, "quest": quest, "user": user})


@app.post("/admin/edit/{quest_id}")
def edit_post(request: Request,
              quest_id: int,
              title: str = Form(...),
              description: str = Form(""),
              organizer_email: str = Form("alibi@mail.ru"),
              price: int = Form(2000),
              genres: List[str] = Form(...),
              difficulty: str = Form(""),
              fear_level: int = Form(0),
              players: int = Form(1),
              image: Optional[UploadFile] = File(None),
              clipboard_image: str = Form(None),
              db: Session = Depends(get_db),
              user=Depends(require_admin)):
    quest = crud.get_quest(db, quest_id)
    if not quest:
        raise HTTPException(404, "Квест не найден")

    # Обработка изображения из буфера обмена
    if clipboard_image and clipboard_image.startswith('data:image'):
        import base64
        # Удаляем старое изображение если есть
        if quest.image_path:
            old_path = os.path.join("static", quest.image_path)
            if os.path.exists(old_path):
                os.remove(old_path)

        # Извлекаем данные из Data URL
        image_data = clipboard_image.split(',')[1]
        image_bytes = base64.b64decode(image_data)

        # Сохраняем файл
        ext = '.png'
        if 'image/jpeg' in clipboard_image:
            ext = '.jpg'
        elif 'image/png' in clipboard_image:
            ext = '.png'
        elif 'image/gif' in clipboard_image:
            ext = '.gif'

        safe_name = f"{uuid.uuid4().hex}{ext}"
        dest = os.path.join("static", "uploads", safe_name)

        with open(dest, "wb") as f:
            f.write(image_bytes)

        quest.image_path = f"uploads/{safe_name}"

    # Обработка обычной загрузки файла
    elif image and image.filename:
        if quest.image_path:
            old_path = os.path.join("static", quest.image_path)
            if os.path.exists(old_path):
                os.remove(old_path)
        image_path = save_upload(image)
        quest.image_path = image_path

    genre_str = ", ".join(genres)

    quest.title = title
    quest.description = description
    quest.organizer_email = organizer_email
    quest.price = price
    quest.genre = genre_str
    quest.difficulty = difficulty
    quest.fear_level = fear_level
    quest.players = players

    db.commit()
    return RedirectResponse("/admin", status_code=303)


@app.post("/admin/delete/{quest_id}")
def admin_delete(request: Request, quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    # Проверяем, есть ли бронирования у этого квеста
    has_bookings = crud.has_quest_bookings(db, quest_id)

    if has_bookings:
        # Получаем информацию о квесте и его бронированиях
        quest = crud.get_quest(db, quest_id)
        quest_bookings = crud.get_quest_bookings(db, quest_id)
        quests = crud.get_quests(db, skip=0, limit=1000, filters={})

        return templates.TemplateResponse("admin_dashboard.html", {
            "request": request,
            "quests": quests,
            "user": user,
            "error": f"Невозможно удалить квест '{quest.title}': есть активные бронирования ({len(quest_bookings)} шт.).",
            "blocked_quest_id": quest_id,
            "quest_bookings": quest_bookings
        })

    # Если бронирований нет, удаляем квест
    q = crud.get_quest(db, quest_id)
    if q and q.image_path:
        file_path = os.path.join("static", q.image_path)
        if os.path.exists(file_path):
            os.remove(file_path)
    crud.delete_quest(db, quest_id)
    return RedirectResponse("/admin", status_code=303)


@app.get("/admin/bookings", response_class=HTMLResponse)
def admin_bookings(request: Request, quest_id: Optional[int] = None, db: Session = Depends(get_db),
                   user=Depends(require_admin)):
    """Страница управления бронированиями для администратора"""
    if quest_id:
        bookings = crud.get_quest_bookings(db, quest_id)
    else:
        bookings = crud.get_all_bookings(db)

    # Получаем все квесты для фильтра
    quests = crud.get_quests(db, skip=0, limit=1000, filters={})

    return templates.TemplateResponse("admin_bookings.html", {
        "request": request,
        "bookings": bookings,
        "quests": quests,
        "user": user,
        "now": datetime.now
    })


@app.post("/admin/delete-quest-with-bookings/{quest_id}")
def admin_delete_quest_with_bookings(quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    """Удаляет квест вместе со всеми его бронированиями"""
    # Сначала удаляем все бронирования квеста
    crud.delete_quest_bookings(db, quest_id)

    # Затем удаляем сам квест
    q = crud.get_quest(db, quest_id)
    if q and q.image_path:
        file_path = os.path.join("static", q.image_path)
        if os.path.exists(file_path):
            os.remove(file_path)
    crud.delete_quest(db, quest_id)

    return RedirectResponse("/admin", status_code=303)


@app.post("/admin/delete-all-bookings/{quest_id}")
def admin_delete_all_bookings(quest_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    """Удаляет все бронирования квеста без удаления самого квеста"""
    crud.delete_quest_bookings(db, quest_id)
    return RedirectResponse("/admin", status_code=303)


@app.post("/admin/delete-booking/{booking_id}")
def admin_delete_booking(booking_id: int, db: Session = Depends(get_db), user=Depends(require_admin)):
    """Удаление бронирования администратором"""
    success = crud.delete_booking(db, booking_id)
    if not success:
        return JSONResponse({"success": False, "message": "Бронирование не найдено"}, status_code=404)

    return RedirectResponse("/admin/bookings", status_code=303)


@app.get("/api/quest-has-bookings/{quest_id}")
def api_quest_has_bookings(quest_id: int, db: Session = Depends(get_db)):
    """API для проверки наличия бронирований у квеста"""
    has_bookings = crud.has_quest_bookings(db, quest_id)
    return JSONResponse({"has_bookings": has_bookings})


# --- Отчеты ---
@app.get("/admin/report/excel")
async def report_excel(db: Session = Depends(get_db), user=Depends(require_admin)):
    """Генерация отчета в Excel с логотипом и печатью"""
    bookings = crud.get_all_bookings(db)

    wb = Workbook()
    ws = wb.active
    ws.title = "Отчет по бронированиям"

    # Стили
    header_font = Font(bold=True, size=16)
    title_font = Font(bold=True, size=12)
    normal_font = Font(size=10)
    bold_font = Font(bold=True, size=10)

    center_align = Alignment(horizontal='center', vertical='center')
    left_align = Alignment(horizontal='left', vertical='center')

    # Добавляем логотип
    try:
        logo_path = "static/images/logo_black.png"
        if os.path.exists(logo_path):
            from openpyxl.drawing.image import Image as XLImage
            logo = XLImage(logo_path)
            logo.width = 80
            logo.height = 80
            ws.add_image(logo, 'A1')
    except:
        pass

    # Шапка документа (смещаем из-за логотипа)
    ws.merge_cells('D1:F1')
    ws['D1'] = "Алиби"
    ws['D1'].font = Font(bold=True, size=18)
    ws['D1'].alignment = center_align

    ws.merge_cells('D2:F2')
    ws['D2'] = "РОССИЯ, 125009, г.Москва, ул.Квестовая, д.88"
    ws['D2'].font = normal_font
    ws['D2'].alignment = center_align

    ws.merge_cells('D3:F3')
    ws['D3'] = "Телефон: +7(999) 999-99-99"
    ws['D3'].font = normal_font
    ws['D3'].alignment = center_align

    ws.merge_cells('D4:F4')
    ws['D4'] = "e-mail: alibi@mail.ru"
    ws['D4'].font = normal_font
    ws['D4'].alignment = center_align

    # Пустая строка
    ws.row_dimensions[5].height = 15

    # Заголовок отчета
    ws.merge_cells('A6:F6')
    ws['A6'] = "ОТЧЕТ ПО БРОНИРОВАНИЯМ"
    ws['A6'].font = Font(bold=True, size=14)
    ws['A6'].alignment = center_align

    # Информация о документе
    ws.merge_cells('A7:F7')
    ws['A7'] = f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}"
    ws['A7'].font = normal_font
    ws['A7'].alignment = center_align

    # Пустая строка
    ws.row_dimensions[9].height = 15

    # Заголовки таблицы
    headers = ['№', 'Пользователь', 'Email пользователя', 'Название квеста', 'Email организатора', 'Цена (руб)']
    for col, header in enumerate(headers, 1):
        cell = ws.cell(row=10, column=col, value=header)
        cell.font = bold_font
        cell.alignment = center_align
        cell.fill = PatternFill(start_color="E6E6FA", end_color="E6E6FA", fill_type="solid")
        cell.border = Border(left=Side(style='thin'), right=Side(style='thin'),
                             top=Side(style='thin'), bottom=Side(style='thin'))

    # Данные бронирований
    total_revenue = 0
    for row, booking in enumerate(bookings, 11):
        ws.cell(row=row, column=1, value=row - 10).alignment = center_align
        ws.cell(row=row, column=2, value=booking.user.username).alignment = left_align
        ws.cell(row=row, column=3, value=booking.user.email or 'Не указан').alignment = left_align
        ws.cell(row=row, column=4, value=booking.quest.title).alignment = left_align
        ws.cell(row=row, column=5, value=booking.quest.organizer_email).alignment = left_align
        ws.cell(row=row, column=6, value=booking.quest.price).alignment = center_align

        # Добавляем границы для всех ячеек
        for col in range(1, 7):
            ws.cell(row=row, column=col).border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                                         top=Side(style='thin'), bottom=Side(style='thin'))

        total_revenue += booking.quest.price

    # Итоговая строка
    last_row = len(bookings) + 11
    ws.merge_cells(f'A{last_row}:E{last_row}')
    ws[f'A{last_row}'] = "ИТОГО:"
    ws[f'A{last_row}'].font = bold_font
    ws[f'A{last_row}'].alignment = Alignment(horizontal='right', vertical='center')
    ws[f'A{last_row}'].fill = PatternFill(start_color="FFFFE0", end_color="FFFFE0", fill_type="solid")

    ws[f'F{last_row}'] = f"{total_revenue} руб"
    ws[f'F{last_row}'].font = bold_font
    ws[f'F{last_row}'].alignment = center_align
    ws[f'F{last_row}'].fill = PatternFill(start_color="FFFFE0", end_color="FFFFE0", fill_type="solid")

    # Добавляем границы для итоговой строки
    for col in range(1, 7):
        ws.cell(row=last_row, column=col).border = Border(left=Side(style='thin'), right=Side(style='thin'),
                                                          top=Side(style='thin'), bottom=Side(style='thin'))

    # Статистика
    stats_row = last_row + 2
    ws.merge_cells(f'A{stats_row}:F{stats_row}')
    ws[f'A{stats_row}'] = f"Всего бронирований: {len(bookings)} | Общая выручка: {total_revenue} руб"
    ws[f'A{stats_row}'].font = bold_font
    ws[f'A{stats_row}'].alignment = center_align

    # Добавляем печать
    try:
        stamp_path = "static/images/stamp.png"
        if os.path.exists(stamp_path):
            from openpyxl.drawing.image import Image as XLImage
            stamp = XLImage(stamp_path)
            stamp.width = 80
            stamp.height = 80
            # Размещаем печать справа внизу
            stamp_cell = f'F{stats_row + 4}'
            ws.add_image(stamp, stamp_cell)
    except:
        pass

    # Место для подписи
    sign_row = stats_row + 6
    ws.merge_cells(f'A{sign_row}:F{sign_row}')
    ws[f'A{sign_row}'] = "_________________________"
    ws[f'A{sign_row}'].alignment = center_align

    ws.merge_cells(f'A{sign_row + 1}:F{sign_row + 1}')
    ws[f'A{sign_row + 1}'] = "Подпись ответственного лица"
    ws[f'A{sign_row + 1}'].alignment = center_align
    ws[f'A{sign_row + 1}'].font = normal_font

    # Настройка ширины колонок
    column_widths = [8, 20, 25, 30, 25, 15]
    for i, width in enumerate(column_widths, 1):
        ws.column_dimensions[get_column_letter(i)].width = width

    # Настройка высоты строк
    for row in range(1, sign_row + 2):
        if row in [1, 6, 10]:
            ws.row_dimensions[row].height = 25
        else:
            ws.row_dimensions[row].height = 18

    # Сохраняем в буфер
    buffer = io.BytesIO()
    wb.save(buffer)
    buffer.seek(0)

    filename = f"otchet_bronirovaniya_{datetime.now().strftime('%Y%m%d_%H%M')}.xlsx"
    encoded_filename = urllib.parse.quote(filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@app.get("/admin/report/pdf")
async def report_pdf(db: Session = Depends(get_db), user=Depends(require_admin)):
    """Генерация отчета в PDF с поддержкой кириллицы"""
    bookings = crud.get_all_bookings(db)

    buffer = io.BytesIO()

    # Используем ReportLab с правильной кодировкой
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    # Регистрируем шрифт с поддержкой кириллицы (если есть)
    try:
        # Попробуем найти стандартные шрифты с кириллицей
        pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
        pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
        font_name = 'Arial'
    except:
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
            font_name = 'DejaVuSans'
        except:
            font_name = 'Helvetica'  # Базовый шрифт

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=60,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    # Создаем кастомные стили для русского текста
    styles.add(ParagraphStyle(
        name='Russian',
        fontName=font_name,
        fontSize=10,
        leading=12,
    ))
    styles.add(ParagraphStyle(
        name='RussianBold',
        fontName=f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold',
        fontSize=10,
        leading=12,
    ))
    styles.add(ParagraphStyle(
        name='RussianTitle',
        fontName=f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold',
        fontSize=16,
        leading=18,
        alignment=1,  # center
    ))
    styles.add(ParagraphStyle(
        name='RussianHeading',
        fontName=f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold',
        fontSize=14,
        leading=16,
        alignment=1,  # center
    ))

    story = []

    # Добавляем логотип
    try:
        logo_path = "static/images/logo_black.png"
        if os.path.exists(logo_path):
            from reportlab.platypus import Image
            logo = Image(logo_path, width=80, height=80)
            logo.hAlign = 'LEFT'
            story.append(logo)
            story.append(Spacer(1, 10))
    except:
        pass

    # Шапка документа
    story.append(Paragraph("Алиби", styles['RussianTitle']))
    story.append(Paragraph("РОССИЯ, 125009, г.Москва, ул.Квестовая, д.88", styles['Russian']))
    story.append(Paragraph("Телефон: +7(999) 999-99-99", styles['Russian']))
    story.append(Paragraph("e-mail: alibi@mail.ru", styles['Russian']))
    story.append(Spacer(1, 12))

    # Заголовок отчета
    story.append(Paragraph("ОТЧЕТ ПО БРОНИРОВАНИЯМ", styles['RussianHeading']))
    story.append(Paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['Russian']))
    story.append(Spacer(1, 20))

    # Таблица с данными
    if bookings:
        # Заголовки таблицы
        data = [['№', 'Пользователь', 'Квест', 'Цена (руб)']]

        total_revenue = 0
        for i, booking in enumerate(bookings, 1):
            data.append([
                str(i),
                booking.user.username or 'Не указан',
                booking.quest.title,
                f"{booking.quest.price} руб"
            ])
            total_revenue += booking.quest.price

        # Создаем таблицу
        table = Table(data, colWidths=[30, 120, 200, 60])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#E6E6FA')),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('ALIGN', (1, 1), (2, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (-1, 0), f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, 0), 10),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.white),
            ('FONTNAME', (0, 1), (-1, -1), font_name),
            ('FONTSIZE', (0, 1), (-1, -1), 9),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))

        story.append(table)
        story.append(Spacer(1, 20))

        # Итоги
        story.append(Paragraph(f"Всего бронирований: {len(bookings)}", styles['RussianBold']))
        story.append(Paragraph(f"Общая выручка: {total_revenue} руб", styles['RussianBold']))
    else:
        story.append(Paragraph("Нет данных о бронированиях", styles['Russian']))

    story.append(Spacer(1, 30))

    # Добавляем печать
    try:
        stamp_path = "static/images/stamp.png"
        if os.path.exists(stamp_path):
            from reportlab.platypus import Image
            stamp = Image(stamp_path, width=80, height=80)
            stamp.hAlign = 'RIGHT'
            story.append(stamp)
    except:
        pass

    # Подпись
    story.append(Spacer(1, 10))
    story.append(Paragraph("_________________________", styles['Russian']))
    story.append(Paragraph("Подпись ответственного лица", styles['Russian']))

    doc.build(story)
    buffer.seek(0)

    filename = f"otchet_bronirovaniya_{datetime.now().strftime('%Y%m%d_%H%M')}.pdf"
    encoded_filename = urllib.parse.quote(filename)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@app.get("/admin/report/word")
async def report_word(db: Session = Depends(get_db), user=Depends(require_admin)):
    """Генерация отчета в Word с логотипом и печатью"""
    try:
        bookings = crud.get_all_bookings(db)

        doc = Document()

        # Настройка стилей
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Arial'
        font.size = Pt(10)

        # Создаем таблицу для шапки с логотипом
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        header_table.columns[0].width = Inches(1.5)
        header_table.columns[1].width = Inches(4.5)

        # Добавляем логотип в первую ячейку
        try:
            logo_path = "static/images/logo_black.png"
            if os.path.exists(logo_path):
                logo_cell = header_table.cell(0, 0)
                logo_paragraph = logo_cell.paragraphs[0]
                logo_run = logo_paragraph.add_run()
                logo_run.add_picture(logo_path, width=Inches(1.8), height=Inches(1.8))
        except:
            pass

        # Добавляем информацию во вторую ячейку
        info_cell = header_table.cell(0, 1)
        info_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_cell.paragraphs[0].add_run("Алиби\n").bold = True
        info_cell.paragraphs[0].add_run("РОССИЯ, 125009, г.Москва, ул.Квестовая, д.88\n")
        info_cell.paragraphs[0].add_run("Телефон: +7(999) 999-99-99\n")
        info_cell.paragraphs[0].add_run("e-mail: alibi@mail.ru")

        doc.add_paragraph()

        # Заголовок отчета
        report_title = doc.add_paragraph("ОТЧЕТ ПО БРОНИРОВАНИЯМ")
        report_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        report_title.runs[0].bold = True
        report_title.runs[0].font.size = Pt(14)

        # Дата формирования
        date_para = doc.add_paragraph(f"Дата формирования: {datetime.now().strftime('%d.%m.%Y %H:%M')}")
        date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Создаем таблицу с данными
        if bookings:
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            table.autofit = False
            table.columns[0].width = Inches(0.5)   # №
            table.columns[1].width = Inches(1.5)   # Пользователь
            table.columns[2].width = Inches(2.5)   # Квест
            table.columns[3].width = Inches(1.0)   # Цена

            # Заголовки таблицы
            headers = ['№', 'Пользователь', 'Квест', 'Цена (руб)']
            hdr_cells = table.rows[0].cells
            for i, header in enumerate(headers):
                hdr_cells[i].text = header
                hdr_cells[i].paragraphs[0].runs[0].bold = True
                hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                # Заливаем фон заголовков
                shading_elm = parse_xml(r'<w:shd {} w:fill="E6E6FA"/>'.format(nsdecls('w')))
                hdr_cells[i]._tc.get_or_add_tcPr().append(shading_elm)

            # Данные
            total_revenue = 0
            for i, booking in enumerate(bookings, 1):
                row_cells = table.add_row().cells
                row_cells[0].text = str(i)
                row_cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                row_cells[1].text = booking.user.username or 'Не указан'
                row_cells[2].text = booking.quest.title

                row_cells[3].text = str(booking.quest.price)
                row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

                total_revenue += booking.quest.price

            doc.add_paragraph()

            # Итоги
            total_para = doc.add_paragraph()
            total_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            total_para.add_run(f"ИТОГО: {total_revenue} руб\n").bold = True
            total_para.add_run(f"Всего бронирований: {len(bookings)} | Общая выручка: {total_revenue} руб").bold = True

        else:
            doc.add_paragraph("Нет данных о бронированиях")

        doc.add_paragraph()
        doc.add_paragraph()

        # Создаем таблицу для подписи и печати
        footer_table = doc.add_table(rows=1, cols=2)
        footer_table.autofit = False
        footer_table.columns[0].width = Inches(4.0)
        footer_table.columns[1].width = Inches(2.0)

        # Подпись в левой ячейке
        sign_cell = footer_table.cell(0, 0)
        sign_cell.paragraphs[0].add_run("_________________________\n")
        sign_cell.paragraphs[0].add_run("Подпись ответственного лица")

        # Печать в правой ячейке
        try:
            stamp_path = "static/images/stamp.png"
            if os.path.exists(stamp_path):
                stamp_cell = footer_table.cell(0, 1)
                stamp_paragraph = stamp_cell.paragraphs[0]
                stamp_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                stamp_run = stamp_paragraph.add_run()
                stamp_run.add_picture(stamp_path, width=Inches(1.8), height=Inches(1.8))
        except:
            pass

        # Сохраняем в буфер
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f"otchet_bronirovaniya_{datetime.now().strftime('%Y%m%d_%H%M')}.docx"
        encoded_filename = urllib.parse.quote(filename)

        return StreamingResponse(
            buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
        )

    except ImportError:
        return JSONResponse({"message": "Для генерации Word отчетов установите python-docx: pip install python-docx"})


# --- Заявление и чек ---
@app.post("/download-statement")
async def download_statement(request: Request, db: Session = Depends(get_db)):
    """Скачивание заявления об отказе от претензий через шаблон Word"""
    user = get_current_user(request, db)
    data = await request.json()

    # Создаем шаблон если его нет
    template_path = "templates/statement_template.docx"
    if not os.path.exists(template_path):
        create_statement_template()

    # Открываем шаблон
    doc = Document(template_path)

    # Данные для замены
    replacements = {
        '{full_name}': data['full_name'],
        '{passport_series}': data['passport_series'],
        '{passport_number}': data['passport_number'],
        '{current_date}': datetime.now().strftime('%d.%m.%Y'),
        '{quest_title}': data.get('quest_title', '')
    }

    # Заменяем метки в документе
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)

    # Также проверяем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for key, value in replacements.items():
                            if key in run.text:
                                run.text = run.text.replace(key, value)

    # Сохраняем в буфер
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"zayavlenie_{data.get('quest_title', 'quest')}.docx"
    encoded_filename = urllib.parse.quote(filename)

    return StreamingResponse(
        buffer,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


@app.post("/download-receipt")
async def download_receipt(request: Request, db: Session = Depends(get_db)):
    """Скачивание чека с поддержкой кириллицы"""
    user = get_current_user(request, db)
    data = await request.json()

    buffer = io.BytesIO()

    # Настройка шрифтов
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.ttfonts import TTFont

    try:
        pdfmetrics.registerFont(TTFont('Arial', 'arial.ttf'))
        pdfmetrics.registerFont(TTFont('Arial-Bold', 'arialbd.ttf'))
        font_name = 'Arial'
    except:
        try:
            pdfmetrics.registerFont(TTFont('DejaVuSans', 'DejaVuSans.ttf'))
            pdfmetrics.registerFont(TTFont('DejaVuSans-Bold', 'DejaVuSans-Bold.ttf'))
            font_name = 'DejaVuSans'
        except:
            font_name = 'Helvetica'

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )

    styles = getSampleStyleSheet()

    # Создаем стили
    styles.add(ParagraphStyle(
        name='ReceiptTitle',
        fontName=f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold',
        fontSize=16,
        leading=18,
        alignment=1,
    ))
    styles.add(ParagraphStyle(
        name='ReceiptText',
        fontName=font_name,
        fontSize=10,
        leading=12,
    ))
    styles.add(ParagraphStyle(
        name='ReceiptBold',
        fontName=f'{font_name}-Bold' if font_name != 'Helvetica' else 'Helvetica-Bold',
        fontSize=10,
        leading=12,
    ))

    story = []

    # Логотип
    try:
        from reportlab.platypus import Image
        logo_path = "static/images/logo_black.png"
        if os.path.exists(logo_path):
            logo = Image(logo_path, width=80, height=80)
            logo.hAlign = 'LEFT'
            story.append(logo)
            story.append(Spacer(1, 10))
    except:
        pass

    # Шапка чека
    story.append(Paragraph("Алиби", styles['ReceiptTitle']))
    story.append(Paragraph("Квест-проект", styles['ReceiptText']))
    story.append(Spacer(1, 15))

    # Реквизиты
    story.append(Paragraph("Юридический адрес: 125009, г. Москва, ул. Квестовая, д. 88", styles['ReceiptText']))
    story.append(Paragraph("ИНН: 7701234567", styles['ReceiptText']))
    story.append(Paragraph("КПП: 770101001", styles['ReceiptText']))
    story.append(Paragraph("ОГРН: 1234567890123", styles['ReceiptText']))
    story.append(Paragraph("Р/с: 40702810123450123456", styles['ReceiptText']))
    story.append(Paragraph('Банк: ПАО "СБЕРБАНК" г. Москва', styles['ReceiptText']))
    story.append(Paragraph("БИК: 044525225", styles['ReceiptText']))
    story.append(Paragraph("К/с: 30101810400000000225", styles['ReceiptText']))

    story.append(Spacer(1, 15))

    # Линия разделитель (имитация)
    story.append(Paragraph("_" * 80, styles['ReceiptText']))
    story.append(Spacer(1, 15))

    # Информация о заказе
    story.append(Paragraph("КАССОВЫЙ ЧЕК", styles['ReceiptBold']))
    story.append(Spacer(1, 10))

    story.append(Paragraph(f"Заказ: {data['quest_title']}", styles['ReceiptText']))
    story.append(Paragraph(f"Дата: {datetime.now().strftime('%d.%m.%Y %H:%M')}", styles['ReceiptText']))
    story.append(Paragraph(f"Клиент: {user.username}", styles['ReceiptText']))
    story.append(Spacer(1, 10))

    # Сумма
    story.append(Paragraph(f"Сумма: {data['quest_price']} руб.", styles['ReceiptBold']))
    story.append(Spacer(1, 10))

    # НДС
    story.append(Paragraph("В том числе НДС 20%: -", styles['ReceiptText']))
    story.append(Paragraph("Согласно Упрощенной системе налогообложения", styles['ReceiptText']))

    story.append(Spacer(1, 20))

    # Печать
    try:
        from reportlab.platypus import Image
        stamp_path = "static/images/stamp.png"
        if os.path.exists(stamp_path):
            stamp = Image(stamp_path, width=80, height=80)
            stamp.hAlign = 'RIGHT'
            story.append(stamp)
    except:
        pass

    # Подпись
    story.append(Spacer(1, 10))
    story.append(Paragraph("Подпись: _________________", styles['ReceiptText']))

    doc.build(story)
    buffer.seek(0)

    filename = f"chek_{data['quest_title']}.pdf"
    encoded_filename = urllib.parse.quote(filename)

    return StreamingResponse(
        buffer,
        media_type="application/pdf",
        headers={"Content-Disposition": f"attachment; filename*=UTF-8''{encoded_filename}"}
    )


if __name__ == "__main__":
    uvicorn.run(
        "main:app",
        host="127.0.0.1",
        port=5000,
        reload=True
    )